"""DOCX text extractor using python-docx."""

from __future__ import annotations

from io import BytesIO

try:
    import docx as _docx_module
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _extract_docx_content(doc) -> str:
    """Extract text content from a python-docx Document object."""
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                try:
                    level_num = int(level)
                    paragraphs.append(f"{'#' * level_num} {text}")
                except ValueError:
                    paragraphs.append(text)
            else:
                paragraphs.append(text)

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            paragraphs.append("\n".join(table_rows))

    return "\n\n".join(paragraphs)


class DocxExtractor:
    """Extract plain text from DOCX bytes."""

    def can_handle(self, doc_type: str) -> bool:
        return doc_type == "docx"

    def extract(self, data: bytes, source_path: str | None = None) -> str:
        if not _DOCX_AVAILABLE:
            raise ImportError("pip install chunkeymonkey[docx]")

        doc = _docx_module.Document(BytesIO(data))
        return _extract_docx_content(doc)
